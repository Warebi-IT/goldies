import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_color):
    """Sets background color of a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell padding."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_training_doc():
    doc = Document()

    # Define color palette (Goldies brand colors)
    COLOR_PRIMARY = RGBColor(227, 86, 42)      # Citra Orange
    COLOR_INK = RGBColor(27, 27, 27)           # Deep Ink
    COLOR_MUTED = RGBColor(100, 100, 100)      # Muted gray
    COLOR_EMERALD = RGBColor(16, 120, 60)      # Success green
    COLOR_AMBER = RGBColor(180, 100, 20)       # Warning amber
    COLOR_ROSE = RGBColor(180, 40, 50)         # Alert red

    # Configure base styles
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = COLOR_INK

    # ---------------------------------------------------------
    # COVER / HEADER TITLE
    # ---------------------------------------------------------
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(30)
    title_p.paragraph_format.space_after = Pt(6)
    run_title = title_p.add_run("GOLDIES TRAVEL")
    run_title.font.size = Pt(28)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_PRIMARY

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(20)
    run_sub = sub_p.add_run("GUIDE OPÉRATIONNEL & MANUEL DE FORMATION UTILISATEURS")
    run_sub.font.size = Pt(16)
    run_sub.font.bold = True
    run_sub.font.color.rgb = COLOR_INK

    desc_p = doc.add_paragraph()
    desc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc_p.paragraph_format.space_after = Pt(30)
    run_desc = desc_p.add_run("Guide complet de prise en main, règles de gestion, cas d'usages et résolution des incidents pour l'équipe d'administration.")
    run_desc.font.size = Pt(11)
    run_desc.font.italic = True
    run_desc.font.color.rgb = COLOR_MUTED

    doc.add_page_break()

    # ---------------------------------------------------------
    # TABLE OF CONTENTS / SUMMARY
    # ---------------------------------------------------------
    h1 = doc.add_heading("SOMMAIRE DU GUIDE", level=1)
    h1.style.font.color.rgb = COLOR_PRIMARY

    toc_items = [
        ("1. Introduction & Présentation de la Plateforme", "Vision, architecture et rôles"),
        ("2. Sécurité & Authentification à Double Facteur (2FA)", "Connexion, activation MFA et réinitialisation"),
        ("3. Module Voyages & Catalogue (Trips)", "Création, modification, liens Stripe, soft-delete et liste d'attente"),
        ("4. Module Inscriptions & Paiements (Bookings)", "Suivi des clientes, validation des paiements et échéancier J-45"),
        ("5. Module Vigilance Voyageuses & Santé (Duty of Care)", "Traitement des alertes médicales et validation de l'assurance"),
        ("6. Module Traçabilité & Journal d'Audit (Audit Events)", "Consultation des traces et historique immuable"),
        ("7. Module Conformité & Maintenance RGPD", "Durées légales, anonymisation et purge en 1 clic"),
        ("8. Module Galerie Photos & Témoignages", "Publication des albums et valorisation des séjours"),
        ("9. Guide Pratique des Cas d'Usages & Dépannage (Troubleshooting)", "Solutions pas-à-pas aux situations courantes"),
        ("10. Check-list Opérationnelle de l'Administrateur", "Routines quotidiennes, avant-départ et post-séjour"),
    ]

    for title, desc in toc_items:
        p = doc.add_paragraph(style='List Bullet')
        r1 = p.add_run(f"{title} : ")
        r1.bold = True
        r2 = p.add_run(desc)
        r2.font.color.rgb = COLOR_MUTED

    doc.add_paragraph().paragraph_format.space_after = Pt(20)

    # ---------------------------------------------------------
    # SECTION 1: INTRODUCTION
    # ---------------------------------------------------------
    h = doc.add_heading("1. INTRODUCTION & PRÉSENTATION DE LA PLATEFORME", level=1)
    h.style.font.color.rgb = COLOR_PRIMARY

    doc.add_paragraph(
        "Goldies Travel est une plateforme numérique dédiée à l'organisation de voyages immersifs et solidaires en groupe 100% féminin en Afrique (Sénégal, Maroc, Tanzanie, Kenya, etc.)."
    )
    doc.add_paragraph(
        "Ce manuel a pour objectif de former l'ensemble des administrateurs, coordinatrices de séjours et gestionnaires opérationnels. "
        "Il garantit une qualité de service irréprochable, une sécurité sans faille pour les voyageuses et une conformité stricte aux réglementations en vigueur (RGPD, Code du Tourisme, Code de la Consommation)."
    )

    # Callout Box for Roles
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = tbl.cell(0, 0)
    set_cell_background(c, "FFF3EB")
    set_cell_margins(c, 150, 150, 200, 200)
    cp = c.paragraphs[0]
    cr1 = cp.add_run("📌 Rôles et Responsabilités dans l'équipe :\n")
    cr1.bold = True
    cr1.font.color.rgb = COLOR_PRIMARY
    cr2 = cp.add_run(
        "• Coordinatrice Séjour : Gestion des inscriptions, suivi médical (allergies), contrôle des assurances, contacts WhatsApp.\n"
        "• Administrateur Ventes / Finances : Création des voyages, configuration des liens Stripe, validation des paiements et des acomptes.\n"
        "• Responsable Technique & Conformité : Exécution de la maintenance RGPD, revue du journal d'audit, gestion des accès 2FA."
    )
    cr2.font.size = Pt(10)

    doc.add_paragraph().paragraph_format.space_after = Pt(15)

    # ---------------------------------------------------------
    # SECTION 2: 2FA / MFA SETUP
    # ---------------------------------------------------------
    h = doc.add_heading("2. SÉCURITÉ & AUTHENTIFICATION DOUBLE FACTEUR (2FA)", level=1)
    h.style.font.color.rgb = COLOR_PRIMARY

    doc.add_paragraph(
        "Pour protéger les données sensibles de nos voyageuses et la configuration financière du site, l'accès au tableau de bord administrateur (/admin) est STRICTEMENT protégé par une double authentification (MFA/TOTP)."
    )

    doc.add_heading("Procédure de Première Connexion & Activation 2FA", level=2)
    steps_2fa = [
        "Rendez-vous sur l'adresse https://goldies.vercel.app/admin.",
        "Saisissez votre adresse email administrateur et votre mot de passe robuste.",
        "Si le 2FA n'est pas encore configuré, vous serez automatiquement redirigé vers l'écran d'enrôlement (/mfa-setup).",
        "Ouvrez votre application d'authentification sur smartphone (Google Authenticator, Microsoft Authenticator ou Authy).",
        "Scannez le QR Code affiché à l'écran.",
        "Saisissez le code à 6 chiffres généré par l'application pour valider l'association.",
        "Votre compte est désormais sécurisé ! Vous pouvez accéder au tableau de bord.",
    ]
    for i, step in enumerate(steps_2fa, 1):
        p = doc.add_paragraph()
        r = p.add_run(f"Étape {i} : ")
        r.bold = True
        p.add_run(step)

    doc.add_heading("Cas d'Usage & Dépannage 2FA", level=2)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Perte de téléphone ou réinitialisation : ").bold = True
    p.add_run("Contactez un co-administrateur ou l'administrateur système pour réinitialiser le facteur TOTP dans Supabase Auth. L'utilisateur sera invité à scanner un nouveau QR code à sa prochaine connexion.")

    doc.add_paragraph().paragraph_format.space_after = Pt(15)

    # ---------------------------------------------------------
    # SECTION 3: VOYAGES & CATALOGUE
    # ---------------------------------------------------------
    h = doc.add_heading("3. MODULE VOYAGES & CATALOGUE (TRIPS)", level=1)
    h.style.font.color.rgb = COLOR_PRIMARY

    doc.add_paragraph(
        "L'onglet 'Gestion des Voyages' permet de piloter l'ensemble de l'offre commerciale de Goldies Travel."
    )

    doc.add_heading("Comment créer ou modifier un voyage ?", level=2)
    trip_fields = [
        ("Nom du voyage", "Ex: 'Sénégal - Immersion & Solidarité' (Nom clair et évocateur)."),
        ("Destination", "Ex: 'Sénégal', 'Maroc', 'Tanzanie'."),
        ("Prix (€)", "Montant total Toutes Taxes Comprises par participante (Ex: 1350)."),
        ("Dates & Durée", "Dates précises (Ex: '15 – 22 Octobre 2026') et durée (Ex: '8 jours / 7 nuits')."),
        ("Capacité (Places max)", "Nombre maximal de participantes (Ex: 12 places). Permet le calcul automatique du taux de remplissage."),
        ("Tag du voyage", "'Disponible' (ouvert aux ventes), 'Bientôt' (teaser), ou 'Complet' (sold out)."),
        ("Description & Programme", "Texte de présentation et découpage détaillé jour par jour (Jour 1, Jour 2, etc.)."),
        ("Inclusions", "Liste des prestations incluses (hébergement, repas, transports locaux, ateliers)."),
        ("Liens de paiement Stripe", "Lien complet (Klarna 3x/4x) et Lien d'acompte (30%)."),
    ]
    for label, explanation in trip_fields:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{label} : ").bold = True
        p.add_run(explanation)

    doc.add_heading("Règles de Gestion Clés sur les Voyages", level=2)
    
    # Box for Rule 1: Soft delete
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = tbl.cell(0, 0)
    set_cell_background(c, "EBF5FB")
    set_cell_margins(c, 120, 120, 180, 180)
    cp = c.paragraphs[0]
    r1 = cp.add_run("🛡️ Règle Anti-Suppression des Voyages (Soft-Delete) :\n")
    r1.bold = True
    r1.font.color.rgb = RGBColor(20, 90, 160)
    r2 = cp.add_run(
        "Si un voyage compte au moins 1 inscription liée, le système INTERDIT sa suppression définitive pour protéger les données comptables et l'historique de la cliente. "
        "Si vous cliquez sur supprimer, le système désactive et archive automatiquement le voyage (is_active = false) pour le masquer du site public sans détruire les dossiers clients."
    )
    r2.font.size = Pt(10)

    p = doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Box for Rule 2: Waitlist
    tbl2 = doc.add_table(rows=1, cols=1)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    c2 = tbl2.cell(0, 0)
    set_cell_background(c2, "FEF9E7")
    set_cell_margins(c2, 120, 120, 180, 180)
    cp2 = c2.paragraphs[0]
    r3 = cp2.add_run("🎟️ Règle de Capacité & Liste d'Attente Automatique :\n")
    r3.bold = True
    r3.font.color.rgb = COLOR_AMBER
    r4 = cp2.add_run(
        "Dès que le nombre de réservations confirmées atteint la capacité maximale (ou si le tag est mis sur 'Complet'), "
        "le bouton de réservation du site public se transforme instantanément en 'Rejoindre la liste d'attente prioritaire'. "
        "Les clientes peuvent s'inscrire sans payer pour être contactées en cas de désistement."
    )
    r4.font.size = Pt(10)

    doc.add_paragraph().paragraph_format.space_after = Pt(15)

    # ---------------------------------------------------------
    # SECTION 4: INSCRIPTIONS & PAIEMENTS
    # ---------------------------------------------------------
    h = doc.add_heading("4. MODULE INSCRIPTIONS & PAIEMENTS (BOOKINGS)", level=1)
    h.style.font.color.rgb = COLOR_PRIMARY

    doc.add_paragraph(
        "L'onglet 'Inscriptions & Contacts' centralise tous les dossiers de voyageuses enregistrés."
    )

    doc.add_heading("Gestion du Statut de Paiement & Modes", level=2)
    doc.add_paragraph(
        "Chaque dossier affiche son statut de paiement en temps réel :"
    )
    p_modes = [
        ("Payé (Vert)", "L'acompte ou la totalité a été encaissé avec succès via Stripe ou virement."),
        ("Non payé (Orange)", "La voyageuse a soumis le formulaire mais n'a pas encore finalisé son règlement."),
        ("Mode 'Acompte' (Bleu)", "La cliente a choisi de régler l'acompte de 30% lors de l'inscription."),
        ("Mode 'Plusieurs fois' (Violet)", "Paiement échelonné (Klarna 3x ou 4x)."),
        ("Mode 'Totalité' (Vert)", "Paiement intégral comptant dès l'inscription."),
    ]
    for mode, desc in p_modes:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{mode} : ").bold = True
        p.add_run(desc)

    doc.add_heading("Règles Financières & Échéancier Solde à J-45", level=2)
    doc.add_paragraph(
        "• Snapshot de Prix Immuable : Le prix du voyage est enregistré de façon définitive lors de la réservation (price_at_booking). Même si le tarif du voyage augmente plus tard, le tarif contractuel de la cliente reste garanti et visible dans le tableau.\n"
        "• Calcul Automatique du Solde à J-45 : Pour toute formule avec acompte, le tableau affiche automatiquement la date limite de règlement du solde (45 jours avant le départ). La coordinatrice peut ainsi envoyer un rappel ciblé à J-50 et J-46."
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(15)

    # ---------------------------------------------------------
    # SECTION 5: VIGILANCE SANTÉ & ASSURANCE
    # ---------------------------------------------------------
    h = doc.add_heading("5. MODULE VIGILANCE VOYAGEUSES & SANTÉ (DUTY OF CARE)", level=1)
    h.style.font.color.rgb = COLOR_PRIMARY

    doc.add_paragraph(
        "La sécurité physique, médicale et alimentaire de nos participantes est la priorité absolue de Goldies Travel."
    )

    doc.add_heading("Détection Automatique des Alertes Médicales & Allergies", level=2)
    doc.add_paragraph(
        "Le système analyse automatiquement le texte saisi par la participante lors de son inscription. Si des termes à risque sont détectés (arachide, cacahuète, gluten, asthme, diabète, femme enceinte, cardiaque, épilepsie, allergie sévère), le système applique immédiatement les actions suivantes :"
    )
    doc.add_paragraph(
        "1. Un badge rouge clignotant 'Alerte Médicale 🚨' apparaît dans la colonne Santé & Allergies du tableau.\n"
        "2. Le compteur 'Alertes Médicales à traiter' du Tableau de Bord s'incrémente.\n"
        "3. La coordinatrice doit prendre contact avec la voyageuse pour adapter les repas et prévenir les guides locaux.\n"
        "4. Une fois l'adaptation confirmée avec les partenaires sur place, la coordinatrice clique sur le badge pour passer le statut en 'Alerte Traitée ✅'."
    )

    doc.add_heading("Contrôle de Conformité de l'Assurance Rapatriement", level=2)
    doc.add_paragraph(
        "Toutes les voyageuses doivent obligatoirement disposer d'une assurance voyage multirisque / assistance rapatriement.\n"
        "• Si la cliente a coché 'Je vais en faire une' : Le bouton affiche 'Attestation requise ⏳'.\n"
        "• À réception du justificatif par email : La coordinatrice clique sur le bouton pour valider le dossier ('Attestation OK ✅').\n"
        "• Règle d'or : Aucune participante ne doit intégrer le groupe WhatsApp final ou monter dans l'avion sans son badge 'Attestation OK'."
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(15)

    # ---------------------------------------------------------
    # SECTION 6: AUDIT & TRAÇABILITÉ
    # ---------------------------------------------------------
    h = doc.add_heading("6. MODULE TRAÇABILITÉ & JOURNAL D'AUDIT (AUDIT EVENTS)", level=1)
    h.style.font.color.rgb = COLOR_PRIMARY

    doc.add_paragraph(
        "Conformément aux exigences des systèmes de confiance et à la réglementation européenne, toutes les opérations critiques sont enregistrées dans un journal d'audit immuable (append-only)."
    )

    doc.add_heading("Événements Enregistrés Automatiquement", level=2)
    events = [
        ("BOOKING_SUBMITTED", "Création d'une nouvelle réservation par une cliente (avec snapshot de prix et statut médical)."),
        ("PAYMENT_STATUS_UPDATED", "Changement de statut de paiement par un administrateur."),
        ("TRIP_CREATED / TRIP_UPDATED", "Création ou mise à jour d'un voyage."),
        ("TRIP_ARCHIVED_PROTECTED", "Archivage sécurisé automatique d'un voyage lors d'une tentative de suppression."),
        ("TRIP_VISIBILITY_TOGGLED", "Activation ou désactivation de la visibilité publique d'un voyage."),
        ("MEDICAL_ALERT_ACK_TOGGLED", "Prise en compte d'une alerte santé par la coordinatrice."),
        ("INSURANCE_VERIFIED_TOGGLED", "Validation de l'attestation d'assurance."),
        ("GDPR_PURGE_EXECUTED", "Exécution de la maintenance et purge périodique RGPD."),
    ]
    for ev, desc in events:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{ev} : ").bold = True
        p.add_run(desc)

    doc.add_paragraph(
        "Comment consulter l'audit ? Cliquez sur le sous-onglet 'Journal d'Audit' dans Inscriptions & Contacts. "
        "Vous pouvez rechercher un événement par date, adresse email ou identifiant."
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(15)

    # ---------------------------------------------------------
    # SECTION 7: RGPD & MAINTENANCE
    # ---------------------------------------------------------
    h = doc.add_heading("7. MODULE CONFORMITÉ & MAINTENANCE RGPD", level=1)
    h.style.font.color.rgb = COLOR_PRIMARY

    doc.add_paragraph(
        "Goldies Travel applique scrupuleusement les durées légales de conservation imposées par le RGPD et le droit commercial français :"
    )

    tbl_gdpr = doc.add_table(rows=5, cols=3)
    tbl_gdpr.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Catégorie de Données", "Durée de Conservation", "Action à Échéance"]
    for j, text in enumerate(headers):
        cell = tbl_gdpr.cell(0, j)
        set_cell_background(cell, "1B1B1B")
        set_cell_margins(cell, 100, 100, 120, 120)
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(10)

    gdpr_rows = [
        ("Messages de contact (Prospects)", "3 ans max (Art. 5 RGPD)", "Suppression définitive automatique"),
        ("Réservations abandonnées / impayées", "90 jours", "Suppression définitive automatique"),
        ("Données de santé & allergies", "1 an max post-séjour", "Purge du champ allergies = '[PURGÉ]'"),
        ("Factures & contrats de vente", "10 ans (Art. L. 123-22 C. Com.)", "Anonymisation irréversible post-10 ans"),
    ]
    for i, row_data in enumerate(gdpr_rows, 1):
        for j, text in enumerate(row_data):
            cell = tbl_gdpr.cell(i, j)
            set_cell_background(cell, "F9F9F9" if i % 2 == 0 else "FFFFFF")
            set_cell_margins(cell, 80, 80, 100, 100)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    doc.add_heading("Procédure de Maintenance RGPD en 1 Clic", level=2)
    doc.add_paragraph(
        "L'administrateur peut à tout moment exécuter la maintenance légale en cliquant sur le bouton 'Maintenance RGPD' en haut à droite de l'onglet Inscriptions. "
        "Le système exécute la procédure stockée sécurisée et affiche un toast récapitulatif du nombre de données purgées."
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(15)

    # ---------------------------------------------------------
    # SECTION 8: GALERIE PHOTOS
    # ---------------------------------------------------------
    h = doc.add_heading("8. MODULE GALERIE PHOTOS & TÉMOIGNAGES", level=1)
    h.style.font.color.rgb = COLOR_PRIMARY

    doc.add_paragraph(
        "L'onglet 'Galerie & Albums' permet de valoriser les retours d'expérience et d'alimenter la preuve sociale du site."
    )
    doc.add_paragraph(
        "• Création d'Album : Titre, destination et photo de couverture.\n"
        "• Ajout de photos : Téléchargement multiple direct vers le stockage cloud sécurisé.\n"
        "• Publication : Switch 'Publié' pour rendre l'album visible dans la section Galerie du site public.\n"
        "• Règle Droit à l'Image : Seules les photos des participantes ayant expressément coché l'autorisation de droit à l'image lors de l'inscription peuvent être publiées."
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(15)

    # ---------------------------------------------------------
    # SECTION 9: TROUBLESHOOTING & CAS D'USAGES
    # ---------------------------------------------------------
    h = doc.add_heading("9. GUIDE PRATIQUE DES CAS D'USAGES & DÉPANNAGE (FAQ)", level=1)
    h.style.font.color.rgb = COLOR_PRIMARY

    doc.add_paragraph(
        "Ce tableau recense les situations opérationnelles les plus fréquentes et la démarche à suivre :"
    )

    trouble_table = doc.add_table(rows=7, cols=3)
    trouble_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_heads = ["Cas d'usage / Problème", "Cause Probable", "Procédure de Résolution Pas-à-Pas"]
    for j, text in enumerate(tbl_heads):
        cell = trouble_table.cell(0, j)
        set_cell_background(cell, "E3562A")
        set_cell_margins(cell, 100, 100, 120, 120)
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(10)

    faq_data = [
        (
            "Une cliente a payé son acompte par virement bancaire au lieu de Stripe.",
            "Paiement hors ligne reçu sur le compte bancaire de la société.",
            "1. Rendez-vous dans 'Inscriptions'.\n2. Recherchez le nom de la cliente.\n3. Cliquez sur le bouton 'Non payé' pour le basculer en 'Payé'.\n4. Vérifiez que le mode est sur 'Acompte'."
        ),
        (
            "Impossible de supprimer définitivement un voyage test.",
            "Des inscriptions de test sont rattachées à ce voyage.",
            "1. Si ce sont de vraies clientes, ne supprimez pas : le voyage sera simplement archivé.\n2. Si c'est un test : supprimez d'abord les réservations de test dans l'onglet Inscriptions, puis supprimez le voyage."
        ),
        (
            "Un voyage complet continue d'afficher 'Réserver'.",
            "La capacité max (qte) n'a pas été renseignée ou le tag est sur 'Disponible'.",
            "1. Modifiez le voyage dans 'Gestion des Voyages'.\n2. Renseignez le nombre de places (ex: 12).\n3. Mettez le Tag sur 'Complet'. Le bouton basculera en liste d'attente."
        ),
        (
            "Une cliente signale une allergie sévère par email après réservation.",
            "Non renseigné lors du formulaire initial.",
            "1. Informez immédiatement la coordinatrice logistique.\n2. Transmettez la consigne au cuisinier/guide local.\n3. Conservez une trace écrite dans le dossier de voyage."
        ),
        (
            "Une cliente demande l'exercice de son droit à l'oubli (RGPD).",
            "Demande officielle par email à contact@goldiestravel.com.",
            "1. Si la cliente n'a jamais voyagé (prospect) : supprimez son message dans Contacts.\n2. Si elle a voyagé : supprimez ses données médicales et conservez uniquement la facture comptable archivée pendant 10 ans."
        ),
        (
            "Le lien de paiement Stripe est refusé ou invalide lors de l'enregistrement.",
            "L'URL ne commence pas par un domaine de paiement sécurisé officiel.",
            "Vérifiez que l'URL commence obligatoirement par https://buy.stripe.com/ ou https://checkout.stripe.com/ pour prévenir les erreurs de saisie."
        ),
    ]

    for i, (cas, cause, sol) in enumerate(faq_data, 1):
        for j, text in enumerate([cas, cause, sol]):
            cell = trouble_table.cell(i, j)
            set_cell_background(cell, "FBFBFB" if i % 2 == 0 else "FFFFFF")
            set_cell_margins(cell, 80, 80, 100, 100)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.size = Pt(9)

    doc.add_paragraph().paragraph_format.space_after = Pt(15)

    # ---------------------------------------------------------
    # SECTION 10: CHECK-LIST QUOTIDIENNE
    # ---------------------------------------------------------
    h = doc.add_heading("10. CHECK-LIST OPÉRATIONNELLE DE L'ADMINISTRATEUR", level=1)
    h.style.font.color.rgb = COLOR_PRIMARY

    doc.add_paragraph(
        "Pour assurer une rigueur constante, appliquez cette grille de contrôle à intervalles réguliers :"
    )

    doc.add_heading("🌅 Routine Quotidienne (5 minutes chaque matin)", level=2)
    daily = [
        "Consulter le Tableau de bord : vérifier les nouveaux paiements et le chiffre d'affaires encaissé.",
        "Traiter les nouveaux messages de contact et répondre sous 24h ouvrées.",
        "Vérifier le bloc 'Vigilance Voyageuses' : traiter les nouvelles alertes médicales.",
    ]
    for d in daily:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(d)

    doc.add_heading("🗓️ Routine Avant-Départ (J-45, J-30, J-7)", level=2)
    pre_trip = [
        "À J-45 : Vérifier les soldes restants à payer pour les clientes en formule Acompte et envoyer le lien de règlement.",
        "À J-30 : Exiger et valider les attestations d'assurance assistance/rapatriement (passer le badge à 'Attestation OK').",
        "À J-15 : Briefing logistique avec les hébergements et cuisiniers locaux sur la liste exacte des allergies traitées.",
        "À J-7 : Création et ouverture du groupe WhatsApp sécurisé réunissant toutes les participantes confirmées.",
    ]
    for pt in pre_trip:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(pt)

    doc.add_heading("🧹 Routine Mensuelle de Clôture & RGPD", level=2)
    monthly = [
        "Cliquer sur le bouton 'Maintenance RGPD' pour purger les anciens prospects et réservations abandonnées.",
        "Vérifier le Journal d'Audit pour s'assurer de l'absence d'actions anormales.",
        "Mettre en ligne les nouveaux albums photos des séjours achevés.",
    ]
    for m in monthly:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(m)

    # Conclusion paragraph
    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    p_end = doc.add_paragraph()
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_end = p_end.add_run("--- Fin du Document de Formation Goldies Travel ---")
    r_end.font.italic = True
    r_end.font.color.rgb = COLOR_MUTED

    # Save document
    output_path = "/Users/yancoubadiatta/Downloads/goldies/Goldies_Guide_Formation_Utilisateurs.docx"
    doc.save(output_path)
    print(f"Document Word généré avec succès : {output_path}")

if __name__ == "__main__":
    create_training_doc()
